"""Ingestor de documentos (.docx, .epub) → JSON transcripts unificados."""

from __future__ import annotations

import html
import json
import re
from datetime import datetime, timezone
from pathlib import Path

import typer
from rich.console import Console
from rich.progress import track

from coach_brain.settings import REPO_ROOT, settings

console = Console()

DOCX_EXT = {".docx"}
EPUB_EXT = {".epub"}

_TAG_RE = re.compile(r"<[^>]+>")
_WS_RE = re.compile(r"[ \t\f\v]+")


def _relpath(p: Path) -> str:
    try:
        return str(p.relative_to(REPO_ROOT))
    except ValueError:
        return str(p)


def _clean(text: str) -> str:
    return _WS_RE.sub(" ", text).strip()


def extract_docx(path: Path) -> list[str]:
    """Devuelve párrafos no vacíos de un .docx (incluye celdas de tablas)."""
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = _clean(para.text)
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [_clean(c.text) for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return parts


def extract_epub(path: Path) -> list[str]:
    """Devuelve bloques de texto (uno por capítulo) de un .epub."""
    import ebooklib
    from ebooklib import epub

    book = epub.read_epub(str(path))
    parts: list[str] = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        raw = item.get_content().decode("utf-8", errors="ignore")
        text = html.unescape(_TAG_RE.sub(" ", raw))
        text = "\n".join(_clean(line) for line in text.splitlines())
        text = re.sub(r"\n{2,}", "\n\n", text).strip()
        if text:
            parts.append(text)
    return parts


def build_record(path: Path, source_type: str, parts: list[str]) -> dict:
    segments = [
        {"id": i, "text": t, "start": None, "end": None, "page": None}
        for i, t in enumerate(parts)
    ]
    return {
        "source_id": f"{source_type}_{path.stem}",
        "source_type": source_type,
        "source_path": _relpath(path),
        "title": path.stem,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "language": settings.whisper_language,
        "metadata": {"pages": None, "duration_seconds": None, "blocks": len(parts)},
        "segments": segments,
        "full_text": "\n\n".join(parts),
    }


def _collect_targets() -> list[tuple[Path, str]]:
    base = settings.raw_path / "docs"
    if not base.exists():
        return []
    targets: list[tuple[Path, str]] = []
    for p in base.rglob("*"):
        suf = p.suffix.lower()
        if suf in DOCX_EXT:
            targets.append((p, "docx"))
        elif suf in EPUB_EXT:
            targets.append((p, "epub"))
    return sorted(targets, key=lambda t: t[0].name)


def run(force: bool = False) -> None:
    targets = _collect_targets()
    if not targets:
        console.print(f"[yellow]No hay .docx/.epub en {settings.raw_path / 'docs'}[/yellow]")
        return

    out_dir = settings.transcripts_path
    out_dir.mkdir(parents=True, exist_ok=True)

    console.print(f"[cyan]Procesando {len(targets)} documento(s)...[/cyan]")
    ok = skip = err = empty = 0

    for path, source_type in track(targets, description="Docs"):
        out_path = out_dir / f"{source_type}_{path.stem}.json"
        if out_path.exists() and not force:
            skip += 1
            continue
        try:
            parts = extract_docx(path) if source_type == "docx" else extract_epub(path)
            if not parts:
                console.print(f"[yellow]Sin texto extraíble: {path.name}[/yellow]")
                empty += 1
                continue
            data = build_record(path, source_type, parts)
            out_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
            ok += 1
        except Exception as e:  # noqa: BLE001
            console.print(f"[red]Error en {path.name}: {e}[/red]")
            err += 1

    console.print(
        f"[green]Listo[/green] — ok={ok} skip={skip} sin_texto={empty} errores={err}"
    )


if __name__ == "__main__":
    typer.run(run)
